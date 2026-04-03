import streamlit as st
import google.generativeai as genai
import json
from datetime import datetime
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
import io
import os

# ─────────────────────────────────────────
# 기존 제안서 예시 데이터 (Few-shot 학습용)
# 새 제안서 추가 시 이 리스트에 dict를 추가하고 GitHub에 push하세요
# ─────────────────────────────────────────
PROPOSAL_EXAMPLES = [
    {
        "client": "한국제조(주)",
        "industry": "제조업",
        "target": "중간관리자(과장~차장급)",
        "headcount": "40명",
        "topic": "리더십 역량 강화",
        "format": "집합교육(오프라인)",
        "duration": "2일(16H)",
        "budget": "800만원",
        "needs": "MZ세대 부하직원과의 소통 어려움, 성과관리 역량 부족",
        "sections": {
            "executive_summary": "귀사 중간관리자의 리더십 역량 강화를 통해 조직 성과를 극대화하고 세대 간 소통 문화를 정착시키기 위한 맞춤형 2일 과정을 제안드립니다. 본 과정은 진단-이론-실습-적용의 4단계 구조로 설계되어 현업 복귀 후 즉시 활용 가능한 실천 역량을 배양합니다. 귀사의 제조업 현장 특성을 반영한 사례 중심 커리큘럼으로 구성하였습니다.",
            "needs_analysis": "첫째, MZ세대 직원과의 세대 간 가치관 차이로 인한 소통 단절이 팀 성과에 영향을 미치고 있습니다. 둘째, 목표 설정 및 피드백 역량 부족으로 성과관리 체계가 개인 역량에 의존하는 상황입니다. 본 과정은 이 두 가지 핵심 문제를 집중적으로 해결합니다.",
            "curriculum": [
                "Day 1 오전 | 리더십 스타일 진단 및 자기이해 (4H)",
                "Day 1 오후 | MZ세대 이해와 세대 간 소통 스킬 워크숍 (4H)",
                "Day 2 오전 | 성과관리와 목표설정(OKR) 실습 (4H)",
                "Day 2 오후 | 코칭 피드백 스킬 & 현업 적용 계획 수립 (4H)"
            ],
            "budget_detail": "강사비 500만원 / 교재·진단도구 150만원 / 장소·식음료 150만원",
            "why_us": "국내 500대 기업 120개사 납품 실적 보유, 제조업 특화 전문 강사진 구성, 교육 후 3개월 사후 코칭 무상 제공"
        }
    },
    {
        "client": "미래금융(주)",
        "industry": "금융업",
        "target": "신입사원",
        "headcount": "80명",
        "topic": "조직 적응 및 비즈니스 스킬",
        "format": "혼합형(블렌디드)",
        "duration": "3일(24H)",
        "budget": "1,200만원",
        "needs": "신입사원의 빠른 조직 적응과 금융권 기초 비즈니스 에티켓 습득",
        "sections": {
            "executive_summary": "귀사 신입사원이 금융업의 특수한 고객 서비스 환경에 빠르게 적응하고 전문가로 성장할 수 있도록 온·오프라인을 결합한 3일 블렌디드 과정을 제안드립니다. 사전 온라인 학습으로 기초를 다지고, 집합교육에서 실전 역량을 완성합니다. 자체 LMS 플랫폼을 통해 수료 후에도 지속적인 학습을 지원합니다.",
            "needs_analysis": "금융업은 높은 수준의 고객 응대 역량과 컴플라이언스 이해가 필수입니다. 신입사원이 이를 단기간에 습득하려면 체계적인 온보딩 프로그램이 필요하며, 특히 디지털 금융 환경에 맞는 비즈니스 커뮤니케이션 역량이 중요합니다.",
            "curriculum": [
                "사전학습(온라인) | 금융업 기초·컴플라이언스 이해 (4H, 자율수강)",
                "Day 1 | 조직문화 이해 및 비즈니스 에티켓 실습 (8H)",
                "Day 2 | 고객 응대 스킬 & 디지털 소통 역량 (8H)",
                "Day 3 | 실전 롤플레이 & 팀 프로젝트 발표 (4H)"
            ],
            "budget_detail": "강사비 700만원 / 온라인 콘텐츠 제작 250만원 / 교재·운영비 250만원",
            "why_us": "금융권 온보딩 과정 연간 50회 이상 운영 경험, 자체 LMS 플랫폼 무상 제공, 수료 후 6개월 온라인 복습 콘텐츠 무제한 제공"
        }
    }
]

SYSTEM_PROMPT = f"""당신은 15년 경력의 기업교육 전문 컨설턴트입니다.
아래 [예시 제안서]들의 문체, 구조, 논리 흐름, 설득력 있는 표현 방식을 철저히 학습하고
새로운 고객사 정보에 맞는 맞춤형 제안서를 작성하세요.

[학습할 예시 제안서]
{json.dumps(PROPOSAL_EXAMPLES, ensure_ascii=False, indent=2)}

[작성 원칙]
1. 반드시 JSON 형식으로만 응답 (다른 텍스트, 마크다운 코드블록 금지)
2. 고객사 업종·규모·니즈에 맞는 구체적인 언어 사용
3. executive_summary: 3문장 이상, 핵심 가치 제안 포함
4. needs_analysis: 고객 문제를 2~3가지로 구체화
5. curriculum: 시간 배분 포함, 4~6개 항목
6. budget_detail: 항목별 금액 명시
7. why_us: 차별화 포인트 3가지 이상

[출력 JSON 구조 - 이 형식 그대로 출력]
{{
  "executive_summary": "...",
  "needs_analysis": "...",
  "curriculum": ["항목1", "항목2", "항목3", "항목4"],
  "budget_detail": "...",
  "why_us": "..."
}}"""


def generate_proposal(client_info: dict, api_key: str) -> dict:
    genai.configure(api_key=api_key)
    model = genai.GenerativeModel(
        model_name="gemini-1.5-flash",
        system_instruction=SYSTEM_PROMPT
    )
    user_message = f"""아래 고객사 정보로 제안서를 작성해주세요.

회사명: {client_info['company']}
업종: {client_info['industry']}
교육 대상: {client_info['target']}
교육 인원: {client_info['headcount']}
교육 주제: {client_info['topic']}
교육 형태: {client_info['format']}
교육 기간: {client_info['duration']}
예산: {client_info['budget']}
핵심 니즈: {client_info['needs']}

반드시 JSON 형식으로만 응답하세요. 코드블록(```) 없이 순수 JSON만 출력하세요."""

    response = model.generate_content(user_message)
    raw = response.text.strip()
    raw = raw.replace("```json", "").replace("```", "").strip()
    return json.loads(raw)


def create_docx(client_info: dict, proposal: dict) -> bytes:
    doc = Document()

    title = doc.add_heading("기업교육 제안서", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_heading("기본 정보", level=1)
    table = doc.add_table(rows=5, cols=2)
    table.style = "Table Grid"
    fields = [
        ("고객사", client_info["company"]),
        ("교육 대상", f"{client_info['target']} / {client_info['headcount']}"),
        ("교육 과정", client_info["topic"]),
        ("교육 형태", f"{client_info['format']} / {client_info['duration']}"),
        ("예산", client_info["budget"]),
    ]
    for i, (k, v) in enumerate(fields):
        table.rows[i].cells[0].text = k
        table.rows[i].cells[1].text = v

    doc.add_paragraph()

    sections_map = [
        ("제안 요약", proposal.get("executive_summary", "")),
        ("니즈 분석", proposal.get("needs_analysis", "")),
        ("예산 구성", proposal.get("budget_detail", "")),
        ("선택 이유", proposal.get("why_us", "")),
    ]
    for heading, content in sections_map:
        doc.add_heading(heading, level=1)
        doc.add_paragraph(content)

    doc.add_heading("교육 커리큘럼", level=1)
    for item in proposal.get("curriculum", []):
        doc.add_paragraph(item, style="List Bullet")

    doc.add_paragraph()
    p = doc.add_paragraph(f"작성일: {datetime.today().strftime('%Y년 %m월 %d일')}")
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.getvalue()


# ─────────────────────────────────────────
# Streamlit UI
# ─────────────────────────────────────────
st.set_page_config(
    page_title="기업교육 제안서 생성기",
    page_icon="📋",
    layout="wide"
)

st.title("📋 기업교육 제안서 자동 생성기")
st.caption("AI가 맞춤형 제안서 초안을 자동으로 작성합니다. (Powered by Google Gemini)")

# 사이드바
with st.sidebar:
    st.header("⚙️ 설정")

    # 환경변수에서 먼저 읽고, 없으면 입력창 표시
    env_key = os.environ.get("GEMINI_API_KEY", "")
    if env_key:
        api_key = env_key
        st.success("API Key 연결됨 ✅")
    else:
        api_key = st.text_input(
            "Gemini API Key",
            type="password",
            help="AIza... 형태의 키를 입력하세요\naistudio.google.com에서 무료 발급"
        )

    st.divider()
    st.markdown("**📚 학습된 예시 제안서**")
    for ex in PROPOSAL_EXAMPLES:
        st.markdown(f"- {ex['client']} ({ex['topic']})")
    st.divider()
    st.markdown("**➕ 예시 추가 방법**")
    st.caption("app.py의 PROPOSAL_EXAMPLES 리스트에\n새 dict를 추가하고\nGitHub에 push하면 자동 반영됩니다.")

# 입력 폼
st.subheader("① 고객사 정보 입력")

col1, col2 = st.columns(2)
with col1:
    company   = st.text_input("회사명 *", placeholder="예: (주)삼성전자")
    industry  = st.text_input("업종 *", placeholder="예: IT, 금융, 제조업, 유통")
    target    = st.text_input("교육 대상 *", placeholder="예: 신입사원, 팀장급, 전직원")
    headcount = st.text_input("교육 인원", placeholder="예: 50명")

with col2:
    topic    = st.text_input("교육 주제 *", placeholder="예: 리더십 역량 강화, AI 업무 활용")
    fmt      = st.selectbox("교육 형태", [
        "집합교육(오프라인)", "온라인(비실시간)",
        "화상교육(실시간)", "혼합형(블렌디드)"
    ])
    duration = st.text_input("교육 기간", placeholder="예: 1일(8H), 2박3일")
    budget   = st.text_input("예산", placeholder="예: 500만원, 협의")

needs = st.text_area(
    "핵심 니즈 / 요청사항 *",
    placeholder="예: 조직 내 세대 간 갈등 해소와 협업 문화 구축이 필요합니다. 특히 팀장급의 코칭 역량이 부족한 상황입니다.",
    height=110
)

st.divider()

# 생성 버튼
if st.button("🚀 제안서 생성하기", type="primary", use_container_width=True):
    if not api_key:
        st.error("왼쪽 사이드바에 Gemini API Key를 입력해주세요. (aistudio.google.com에서 무료 발급)")
    elif not all([company, industry, target, topic, needs]):
        st.warning("* 표시된 항목은 필수입니다: 회사명, 업종, 교육 대상, 교육 주제, 핵심 니즈")
    else:
        client_info = {
            "company": company, "industry": industry,
            "target": target, "headcount": headcount or "미정",
            "topic": topic, "format": fmt,
            "duration": duration or "미정", "budget": budget or "협의",
            "needs": needs
        }
        with st.spinner("AI가 제안서를 작성 중입니다... (약 10~20초)"):
            try:
                proposal = generate_proposal(client_info, api_key)
                st.session_state["proposal"] = proposal
                st.session_state["client_info"] = client_info
                st.success("✅ 제안서가 생성되었습니다!")
            except json.JSONDecodeError:
                st.error("AI 응답을 파싱하는 데 실패했습니다. 다시 시도해주세요.")
            except Exception as e:
                st.error(f"오류 발생: {str(e)}")

# 결과 출력
if "proposal" in st.session_state:
    p  = st.session_state["proposal"]
    ci = st.session_state["client_info"]

    st.subheader("② 생성된 제안서 미리보기")

    tab1, tab2, tab3, tab4, tab5 = st.tabs([
        "📌 제안 요약", "🔍 니즈 분석",
        "📅 커리큘럼", "💰 예산 구성", "⭐ 선택 이유"
    ])
    with tab1:
        st.write(p.get("executive_summary", ""))
    with tab2:
        st.write(p.get("needs_analysis", ""))
    with tab3:
        for item in p.get("curriculum", []):
            st.markdown(f"- {item}")
    with tab4:
        st.write(p.get("budget_detail", ""))
    with tab5:
        st.write(p.get("why_us", ""))

    st.divider()
    st.subheader("③ Word 파일 다운로드")

    docx_bytes = create_docx(ci, p)
    filename   = f"제안서_{ci['company']}_{datetime.today().strftime('%Y%m%d')}.docx"

    st.download_button(
        label="📥 Word(.docx) 파일 다운로드",
        data=docx_bytes,
        file_name=filename,
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        use_container_width=True
    )
